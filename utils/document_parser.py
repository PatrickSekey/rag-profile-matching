"""
Document Parser Module
Handles extraction of text from various resume file formats (PDF, TXT, DOCX)
"""

import os
import re
from typing import Dict, Optional
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse different document formats and extract text content"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']
    
    def parse_document(self, file_path: str) -> Dict[str, str]:
        """
        Parse document and extract text based on file extension
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with 'text' key containing extracted content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._parse_pdf_robust(file_path)
        elif file_extension == '.docx':
            return self._parse_docx(file_path)
        elif file_extension == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported: {self.supported_formats}")
    
    def _parse_pdf_robust(self, file_path: str) -> Dict[str, str]:
        """
        Extract text from PDF using multiple methods with fallbacks
        """
        text = ""
        
        # Method 1: Try pymupdf (PyMuPDF)
        try:
            import pymupdf as fitz
            logger.info(f"Method 1: Trying pymupdf for {os.path.basename(file_path)}")
            doc = fitz.open(file_path)
            
            # Handle encrypted PDFs
            if doc.is_encrypted:
                logger.info("PDF is encrypted, trying to decrypt...")
                try:
                    doc.authenticate("")
                except:
                    pass
            
            text = ""
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    logger.warning(f"Could not extract page {page_num}: {str(e)}")
                    continue
            
            doc.close()
            
            if text and len(text.strip()) > 100:
                logger.info(f"✅ pymupdf extracted {len(text)} characters")
                return {"text": text.strip(), "file_type": "pdf", "method": "pymupdf"}
            else:
                logger.warning(f"pymupdf extracted only {len(text)} characters, trying other methods...")
                
        except ImportError:
            logger.warning("pymupdf not installed, trying other methods...")
        except Exception as e:
            logger.warning(f"pymupdf error: {str(e)}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            logger.info(f"Method 2: Trying pdfplumber for {os.path.basename(file_path)}")
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text and len(text.strip()) > 100:
                logger.info(f"✅ pdfplumber extracted {len(text)} characters")
                return {"text": text.strip(), "file_type": "pdf", "method": "pdfplumber"}
            else:
                logger.warning(f"pdfplumber extracted only {len(text)} characters")
                
        except ImportError:
            logger.warning("pdfplumber not installed, trying next method...")
        except Exception as e:
            logger.warning(f"pdfplumber error: {str(e)}")
        
        # Method 3: Try PyPDF2
        try:
            import PyPDF2
            logger.info(f"Method 3: Trying PyPDF2 for {os.path.basename(file_path)}")
            text = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text and len(text.strip()) > 100:
                logger.info(f"✅ PyPDF2 extracted {len(text)} characters")
                return {"text": text.strip(), "file_type": "pdf", "method": "PyPDF2"}
                
        except ImportError:
            logger.warning("PyPDF2 not installed")
        except Exception as e:
            logger.warning(f"PyPDF2 error: {str(e)}")
        
        # Method 4: Try pdfminer
        try:
            from pdfminer.high_level import extract_text
            logger.info(f"Method 4: Trying pdfminer for {os.path.basename(file_path)}")
            text = extract_text(file_path)
            
            if text and len(text.strip()) > 100:
                logger.info(f"✅ pdfminer extracted {len(text)} characters")
                return {"text": text.strip(), "file_type": "pdf", "method": "pdfminer"}
                
        except ImportError:
            logger.warning("pdfminer not installed")
        except Exception as e:
            logger.warning(f"pdfminer error: {str(e)}")
        
        # Method 5: Try tika (Apache Tika) - most reliable but slower
        try:
            from tika import parser
            logger.info(f"Method 5: Trying Tika for {os.path.basename(file_path)}")
            parsed = parser.from_file(file_path)
            text = parsed.get('content', '')
            
            if text and len(text.strip()) > 100:
                logger.info(f"✅ Tika extracted {len(text)} characters")
                return {"text": text.strip(), "file_type": "pdf", "method": "tika"}
                
        except ImportError:
            logger.warning("tika not installed")
        except Exception as e:
            logger.warning(f"Tika error: {str(e)}")
        
        # If all methods fail, try to extract at least some text
        if text and len(text.strip()) > 0:
            logger.warning(f"⚠️ Partial extraction: {len(text)} characters from {os.path.basename(file_path)}")
            return {"text": text.strip(), "file_type": "pdf", "method": "partial"}
        
        # Final fallback: raise error
        raise ValueError(f"Could not extract text from PDF: {file_path}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, str]:
        """Original PDF parser (kept for backward compatibility)"""
        return self._parse_pdf_robust(file_path)
    
    def _parse_docx(self, file_path: str) -> Dict[str, str]:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Clean up extra whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            if text:
                logger.info(f"Successfully parsed DOCX: {os.path.basename(file_path)} - {len(text)} characters")
                return {"text": text, "file_type": "docx", "method": "python-docx"}
            else:
                raise ValueError("No text extracted from DOCX")
                
        except ImportError:
            logger.error("python-docx not installed")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: str) -> Dict[str, str]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Clean up extra whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            logger.info(f"Successfully parsed TXT: {os.path.basename(file_path)} - {len(text)} characters")
            return {"text": text, "file_type": "txt", "method": "txt"}
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                text = re.sub(r'\s+', ' ', text).strip()
                logger.info(f"Successfully parsed TXT (latin-1): {os.path.basename(file_path)} - {len(text)} characters")
                return {"text": text, "file_type": "txt", "method": "txt"}
            except Exception as e:
                logger.error(f"Error parsing TXT {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise
    
    def get_file_metadata(self, file_path: str) -> Dict:
        """Get basic metadata about the file"""
        return {
            "filename": os.path.basename(file_path),
            "file_size": os.path.getsize(file_path),
            "file_extension": os.path.splitext(file_path)[1],
            "file_path": file_path
        }
